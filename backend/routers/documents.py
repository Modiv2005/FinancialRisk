"""Document management router - upload, process, and manage documents."""
import os
import uuid
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.orm import Session
from typing import Optional, List
from database import get_db
from auth import get_current_user
from config import settings
import models
import schemas

router = APIRouter(prefix="/api/documents", tags=["Documents"])

ALLOWED_EXTENSIONS = {".pdf", ".csv", ".xlsx", ".xls", ".docx", ".txt", ".json"}


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


async def process_document_background(doc_id: str, file_path: str, file_type: str):
    """Background task to process uploaded documents."""
    from database import SessionLocal
    db = SessionLocal()
    try:
        doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
        if not doc:
            return
        
        doc.status = "processing"
        db.commit()
        
        content = ""
        try:
            if file_type == ".csv":
                import pandas as pd
                df = pd.read_csv(file_path)
                content = df.to_string()[:10000]
                doc.metadata_json = {
                    "rows": len(df), "columns": list(df.columns),
                    "dtypes": {k: str(v) for k, v in df.dtypes.items()},
                }
            elif file_type in (".xlsx", ".xls"):
                import pandas as pd
                df = pd.read_excel(file_path)
                content = df.to_string()[:10000]
                doc.metadata_json = {
                    "rows": len(df), "columns": list(df.columns),
                    "dtypes": {k: str(v) for k, v in df.dtypes.items()},
                }
            elif file_type == ".pdf":
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        pages_text = [p.extract_text() or "" for p in pdf.pages]
                        content = "\n\n".join(pages_text)[:15000]
                        doc.metadata_json = {"pages": len(pdf.pages)}
                except Exception:
                    content = "[PDF processing unavailable]"
            elif file_type == ".docx":
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(file_path)
                    content = "\n".join([p.text for p in docx_doc.paragraphs])[:15000]
                    doc.metadata_json = {"paragraphs": len(docx_doc.paragraphs)}
                except Exception:
                    content = "[DOCX processing unavailable]"
            elif file_type == ".txt":
                with open(file_path, "r", errors="ignore") as f:
                    content = f.read()[:15000]
            elif file_type == ".json":
                import json
                with open(file_path, "r") as f:
                    data = json.load(f)
                content = json.dumps(data, indent=2)[:15000]
            
            doc.content_text = content
            doc.status = "completed"
            
            # Create chunks for RAG
            if content:
                chunk_size = 1000
                chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
                for idx, chunk_text in enumerate(chunks[:50]):
                    chunk = models.DocumentChunk(
                        document_id=doc.id,
                        chunk_index=idx,
                        content=chunk_text,
                        metadata_json={"source": doc.filename, "chunk": idx},
                    )
                    db.add(chunk)
        except Exception as e:
            doc.status = "failed"
            doc.metadata_json = {"error": str(e)}
        
        db.commit()
    finally:
        db.close()


@router.post("/upload", response_model=schemas.DocumentResponse)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
    user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Upload and process a document."""
    ext = get_file_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Unsupported file type: {ext}")
    
    if file.size and file.size > settings.MAX_UPLOAD_SIZE:
        raise HTTPException(400, "File too large (max 50MB)")
    
    # Save file
    file_id = str(uuid.uuid4())
    file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}{ext}")
    
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    # Create document record
    doc = models.Document(
        id=file_id,
        user_id=user.id,
        filename=file.filename,
        file_type=ext,
        file_size=len(content),
        file_path=file_path,
        status="pending",
        category=category,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    # Process in background
    background_tasks.add_task(process_document_background, doc.id, file_path, ext)
    
    return schemas.DocumentResponse.model_validate(doc)


@router.post("/upload-batch")
async def upload_batch(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Batch upload multiple documents."""
    results = []
    for file in files[:20]:  # Max 20 files per batch
        ext = get_file_extension(file.filename)
        if ext not in ALLOWED_EXTENSIONS:
            results.append({"filename": file.filename, "status": "rejected", "reason": "Unsupported type"})
            continue
        
        file_id = str(uuid.uuid4())
        file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}{ext}")
        
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        doc = models.Document(
            id=file_id, user_id=user.id, filename=file.filename,
            file_type=ext, file_size=len(content), file_path=file_path,
            status="pending",
        )
        db.add(doc)
        background_tasks.add_task(process_document_background, doc.id, file_path, ext)
        results.append({"filename": file.filename, "id": file_id, "status": "queued"})
    
    db.commit()
    return {"uploaded": len(results), "results": results}


@router.get("/", response_model=schemas.DocumentListResponse)
def list_documents(
    skip: int = 0, limit: int = 50,
    status: Optional[str] = None,
    user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List user's documents."""
    query = db.query(models.Document).filter(models.Document.user_id == user.id)
    if status:
        query = query.filter(models.Document.status == status)
    
    total = query.count()
    docs = query.order_by(models.Document.created_at.desc()).offset(skip).limit(limit).all()
    
    return schemas.DocumentListResponse(
        documents=[schemas.DocumentResponse.model_validate(d) for d in docs],
        total=total,
    )


@router.get("/{doc_id}", response_model=schemas.DocumentResponse)
def get_document(doc_id: str, user: models.User = Depends(get_current_user), db: Session = Depends(get_db)):
    """Get document details."""
    doc = db.query(models.Document).filter(
        models.Document.id == doc_id, models.Document.user_id == user.id
    ).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    return schemas.DocumentResponse.model_validate(doc)


@router.delete("/{doc_id}")
def delete_document(doc_id: str, user: models.User = Depends(get_current_user), db: Session = Depends(get_db)):
    """Delete a document."""
    doc = db.query(models.Document).filter(
        models.Document.id == doc_id, models.Document.user_id == user.id
    ).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    
    db.delete(doc)
    db.commit()
    return {"message": "Document deleted"}
